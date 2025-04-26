from docx import Document

# Create a new Word document
doc = Document()

# Title of the document
doc.add_heading('Microsoft Teams Meeting Transcript', 0)

# Meeting details
doc.add_paragraph('Meeting Topic: Economic Sanctions Discussion')
doc.add_paragraph('Date: April 26, 2025')
doc.add_paragraph('Duration: 10 minutes')
doc.add_paragraph('Participants:')
doc.add_paragraph('• John Smith (meherdontoju@gmail.com)')
doc.add_paragraph('• Emily Johnson (mspragai@outlook.com)')
doc.add_paragraph('• Michael Brown (asyamsundar456@gmail.com)')

# Transcript content
transcript = [
    ("[00:00]", "John Smith (meherdontoju@gmail.com)", "Alright, let's kick off the discussion on economic sanctions. I think it’s a crucial topic given the ongoing global tensions. Does anyone want to start with their thoughts on the current situation?"),
    ("[00:15]", "Emily Johnson (mspragai@outlook.com)", "Sure, I can start. So, economic sanctions have become a primary tool in international diplomacy. They’re often used to deter countries from engaging in activities that the international community considers harmful, like human rights violations or nuclear proliferation. But I feel they can have unintended consequences on the civilian population, and that’s something we need to think about more."),
    ("[00:35]", "Michael Brown (asyamsundar456@gmail.com)", "I agree with Emily, but I also think they can be effective in forcing change. Look at how sanctions impacted Iran’s nuclear program or how they’ve affected North Korea. There’s a debate about whether sanctions actually lead to political change or just punish ordinary people."),
    ("[00:55]", "John Smith (meherdontoju@gmail.com)", "Exactly, Michael. And it’s often difficult to measure the success of sanctions. Sometimes they force a country to change its behavior, but other times they just deepen the divide between the government and the people. The key question is, are they truly effective?"),
    ("[01:20]", "Emily Johnson (mspragai@outlook.com)", "And then there's the issue of targeted versus broad-based sanctions. Targeted sanctions—such as freezing the assets of government officials—might be more effective and less harmful to civilians. But broad sanctions, like trade embargoes, tend to have a more widespread impact on the economy."),
    ("[01:45]", "Michael Brown (asyamsundar456@gmail.com)", "Right, targeted sanctions are certainly more precise, but they can also be harder to enforce. And sometimes they don't have the same kind of visible impact, which can make them seem less effective to the public or to decision-makers."),
    ("[02:10]", "John Smith (meherdontoju@gmail.com)", "Another point to consider is the role of allies. Economic sanctions are more likely to be effective if there’s international cooperation. But when countries like Russia or China don’t agree with sanctions, it’s hard to enforce them globally."),
    ("[02:35]", "Emily Johnson (mspragai@outlook.com)", "Exactly. And sanctions can also be undermined by countries that aren't part of the sanctioning coalition. That’s why enforcement becomes a key issue. If a country can still trade with others who aren’t imposing sanctions, it dilutes the pressure."),
    ("[03:00]", "Michael Brown (asyamsundar456@gmail.com)", "There’s also the humanitarian aspect. The problem with sanctions is that they can hurt the people who are already suffering, not necessarily the leadership. For example, when sanctions were imposed on Iraq in the 1990s, it led to widespread poverty and malnutrition. The civilians paid the price, not the regime."),
    ("[03:30]", "John Smith (meherdontoju@gmail.com)", "That’s a really valid concern. It’s always difficult to balance the humanitarian impact with the intended political goal. Ideally, sanctions would be designed to punish those in power without hurting innocent civilians. But that’s not always how it plays out."),
    ("[03:55]", "Emily Johnson (mspragai@outlook.com)", "Absolutely. And that’s why some argue for a more nuanced approach—combining sanctions with diplomatic efforts and humanitarian aid. That way, the sanctions are more focused, and there’s a chance to alleviate the suffering of civilians."),
    ("[04:20]", "Michael Brown (asyamsundar456@gmail.com)", "And I think sanctions should be tied to clear, measurable goals. If a country doesn’t meet those goals, then the sanctions can be escalated. But if the goals are met, the sanctions can be lifted. This kind of flexibility could make sanctions more effective and less harmful in the long run."),
    ("[04:45]", "John Smith (meherdontoju@gmail.com)", "Right, Michael, flexibility could be key to success. It’s about being adaptable. Now, what do you think about the recent sanctions against Russia? Do you think they’ve been effective in changing their behavior?"),
    ("[05:10]", "Emily Johnson (mspragai@outlook.com)", "It’s hard to say. On one hand, the sanctions have certainly put a strain on Russia’s economy. But on the other hand, they’ve also solidified Putin’s support domestically by framing the sanctions as foreign aggression. So, while the economy has suffered, the political situation hasn’t changed that much."),
    ("[05:40]", "Michael Brown (asyamsundar456@gmail.com)", "Yeah, and Russia has shown that it can adapt to sanctions by finding alternative trading partners. So, while the sanctions hurt, they’re not necessarily leading to the desired political changes. This just goes to show that sanctions alone are rarely enough."),
    ("[06:05]", "John Smith (meherdontoju@gmail.com)", "That’s true. It’s a complex issue. I think we need to focus more on multilateral pressure and less on unilateral sanctions. Countries should be working together, not just imposing sanctions on their own."),
    ("[06:30]", "Emily Johnson (mspragai@outlook.com)", "Agreed. And I think another issue is the lack of transparency in how sanctions are applied. Sometimes, it’s unclear what the specific goals are or who the intended targets really are."),
    ("[06:55]", "Michael Brown (asyamsundar456@gmail.com)", "Transparency could definitely improve the effectiveness of sanctions. If the target of the sanctions knows exactly what’s required to have them lifted, they might be more willing to comply. Without that clarity, sanctions risk becoming a blunt instrument that doesn’t achieve much."),
    ("[07:20]", "John Smith (meherdontoju@gmail.com)", "Right, and we need more consistent monitoring to ensure that sanctions are being enforced as intended. Otherwise, they just become symbolic."),
    ("[07:45]", "Emily Johnson (mspragai@outlook.com)", "Speaking of monitoring, I think technology has a big role to play. With the advent of new tools, we can track transactions more closely, ensuring that sanction violations are caught and addressed quickly."),
    ("[08:10]", "Michael Brown (asyamsundar456@gmail.com)", "That’s a great point. Increased monitoring and enforcement through tech could give sanctions a lot more power. But the political will to implement these tools needs to be there."),
    ("[08:30]", "John Smith (meherdontoju@gmail.com)", "Definitely. So, to wrap up, it seems clear that while sanctions have their role, they’re not a perfect tool. We need a more thoughtful, adaptable approach that focuses on targeted actions, transparency, and international cooperation. Anything else before we end?"),
    ("[09:00]", "Emily Johnson (mspragai@outlook.com)", "I think we’ve covered the key points. It’s definitely a complex issue, but I think the discussion has been really productive."),
    ("[09:30]", "Michael Brown (asyamsundar456@gmail.com)", "Agreed. It’s always good to get a broad perspective on such a multifaceted topic. Thanks for the great discussion!"),
    ("[10:00]", "John Smith (meherdontoju@gmail.com)", "Thanks everyone! I’ll follow up with some of the points we discussed today. Talk soon!")
]

# Add transcript to the document
for time, speaker, message in transcript:
    doc.add_paragraph(f'{time}  {speaker}:', style='Heading 3')
    doc.add_paragraph(message)

# Save the document
#doc.save('Economic_Sanctions_Transcript.docx')

#--------------------------------------------------------------------------------

# Json format


import json

# Transcript content in a structured way (similar to Microsoft Teams format)
transcript = [
    {"time": "[00:00]", "speaker": "John Smith (email1@gmail.com)", "message": "Alright, let's kick off the discussion on economic sanctions. I think it’s a crucial topic given the ongoing global tensions. Does anyone want to start with their thoughts on the current situation?"},
    {"time": "[00:15]", "speaker": "Emily Johnson (email2@gmail.com)", "message": "Sure, I can start. So, economic sanctions have become a primary tool in international diplomacy. They’re often used to deter countries from engaging in activities that the international community considers harmful, like human rights violations or nuclear proliferation. But I feel they can have unintended consequences on the civilian population, and that’s something we need to think about more."},
    {"time": "[00:35]", "speaker": "Michael Brown (email3@gmail.com)", "message": "I agree with Emily, but I also think they can be effective in forcing change. Look at how sanctions impacted Iran’s nuclear program or how they’ve affected North Korea. There’s a debate about whether sanctions actually lead to political change or just punish ordinary people."},
    {"time": "[00:55]", "speaker": "John Smith (email1@gmail.com)", "message": "Exactly, Michael. And it’s often difficult to measure the success of sanctions. Sometimes they force a country to change its behavior, but other times they just deepen the divide between the government and the people. The key question is, are they truly effective?"},
    {"time": "[01:20]", "speaker": "Emily Johnson (email2@gmail.com)", "message": "And then there's the issue of targeted versus broad-based sanctions. Targeted sanctions—such as freezing the assets of government officials—might be more effective and less harmful to civilians. But broad sanctions, like trade embargoes, tend to have a more widespread impact on the economy."},
    {"time": "[01:45]", "speaker": "Michael Brown (email3@gmail.com)", "message": "Right, targeted sanctions are certainly more precise, but they can also be harder to enforce. And sometimes they don't have the same kind of visible impact, which can make them seem less effective to the public or to decision-makers."},
    {"time": "[02:10]", "speaker": "John Smith (email1@gmail.com)", "message": "Another point to consider is the role of allies. Economic sanctions are more likely to be effective if there’s international cooperation. But when countries like Russia or China don’t agree with sanctions, it’s hard to enforce them globally."},
    {"time": "[02:35]", "speaker": "Emily Johnson (email2@gmail.com)", "message": "Exactly. And sanctions can also be undermined by countries that aren't part of the sanctioning coalition. That’s why enforcement becomes a key issue. If a country can still trade with others who aren’t imposing sanctions, it dilutes the pressure."},
    {"time": "[03:00]", "speaker": "Michael Brown (email3@gmail.com)", "message": "There’s also the humanitarian aspect. The problem with sanctions is that they can hurt the people who are already suffering, not necessarily the leadership. For example, when sanctions were imposed on Iraq in the 1990s, it led to widespread poverty and malnutrition. The civilians paid the price, not the regime."},
    {"time": "[03:30]", "speaker": "John Smith (email1@gmail.com)", "message": "That’s a really valid concern. It’s always difficult to balance the humanitarian impact with the intended political goal. Ideally, sanctions would be designed to punish those in power without hurting innocent civilians. But that’s not always how it plays out."},
    {"time": "[03:55]", "speaker": "Emily Johnson (email2@gmail.com)", "message": "Absolutely. And that’s why some argue for a more nuanced approach—combining sanctions with diplomatic efforts and humanitarian aid. That way, the sanctions are more focused, and there’s a chance to alleviate the suffering of civilians."},
    {"time": "[04:20]", "speaker": "Michael Brown (email3@gmail.com)", "message": "And I think sanctions should be tied to clear, measurable goals. If a country doesn’t meet those goals, then the sanctions can be escalated. But if the goals are met, the sanctions can be lifted. This kind of flexibility could make sanctions more effective and less harmful in the long run."},
    {"time": "[04:45]", "speaker": "John Smith (email1@gmail.com)", "message": "Right, Michael, flexibility could be key to success. It’s about being adaptable. Now, what do you think about the recent sanctions against Russia? Do you think they’ve been effective in changing their behavior?"},
    {"time": "[05:10]", "speaker": "Emily Johnson (email2@gmail.com)", "message": "It’s hard to say. On one hand, the sanctions have certainly put a strain on Russia’s economy. But on the other hand, they’ve also solidified Putin’s support domestically by framing the sanctions as foreign aggression. So, while the economy has suffered, the political situation hasn’t changed that much."},
    {"time": "[05:40]", "speaker": "Michael Brown (email3@gmail.com)", "message": "Yeah, and Russia has shown that it can adapt to sanctions by finding alternative trading partners. So, while the sanctions hurt, they’re not necessarily leading to the desired political changes. This just goes to show that sanctions alone are rarely enough."},
    {"time": "[06:05]", "speaker": "John Smith (email1@gmail.com)", "message": "That’s true. It’s a complex issue. I think we need to focus more on multilateral pressure and less on unilateral sanctions. Countries should be working together, not just imposing sanctions on their own."},
    {"time": "[06:30]", "speaker": "Emily Johnson (email2@gmail.com)", "message": "Agreed. And I think another issue is the lack of transparency in how sanctions are applied. Sometimes, it’s unclear what the specific goals are or who the intended targets really are."},
    {"time": "[06:55]", "speaker": "Michael Brown (email3@gmail.com)", "message": "Transparency could definitely improve the effectiveness of sanctions. If the target of the sanctions knows exactly what’s required to have them lifted, they might be more willing to comply. Without that clarity, sanctions risk becoming a blunt instrument that doesn’t achieve much."},
    {"time": "[07:20]", "speaker": "John Smith (email1@gmail.com)", "message": "Right, and we need more consistent monitoring to ensure that sanctions are being enforced as intended. Otherwise, they just become symbolic."},
    {"time": "[07:45]", "speaker": "Emily Johnson (email2@gmail.com)", "message": "Speaking of monitoring, I think technology has a big role to play. With the advent of new tools, we can track transactions more closely, ensuring that sanction violations are caught and addressed quickly."},
    {"time": "[08:10]", "speaker": "Michael Brown (email3@gmail.com)", "message": "That’s a great point. Increased monitoring and enforcement through tech could give sanctions a lot more power. But the political will to implement these tools needs to be there."},
    {"time": "[08:30]", "speaker": "John Smith (email1@gmail.com)", "message": "Definitely. So, to wrap up, it seems clear that while sanctions have their role, they’re not a perfect tool. We need a more thoughtful, adaptable approach that focuses on targeted actions, transparency, and international cooperation. Anything else before we end?"},
    {"time": "[09:00]", "speaker": "Emily Johnson (email2@gmail.com)", "message": "I think we’ve covered the key points. It’s definitely a complex issue, but I think the discussion has been really productive."},
    {"time": "[09:30]", "speaker": "Michael Brown (email3@gmail.com)", "message": "Agreed. It’s always good to get a broad perspective on such a multifaceted topic. Thanks for the great discussion!"},
    {"time": "[10:00]", "speaker": "John Smith (email1@gmail.com)", "message": "Thanks everyone! I’ll follow up with some of the points we discussed today. Talk soon!"}
]

# Save transcript to a JSON file
file_path = 'Economic_Sanctions_Transcript.json'
with open(file_path, 'w') as json_file:
    json.dump(transcript, json_file, indent=4)

print(f"Transcript has been saved to {file_path}")
